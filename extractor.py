"""
extractor.py
------------
Handles text extraction from uploaded resume files (PDF or DOCX).

Module 1 responsibility: Extract text content from the uploaded resume
so it can be stored and later analyzed by Module 2 (scoring) and
Module 3 (ATS keyword checking).
"""

import pdfplumber
import docx


def extract_text_from_pdf(filepath):
    """Extract all text from a PDF file, page by page."""
    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def extract_text_from_docx(filepath):
    """Extract all text from a DOCX file, paragraph by paragraph."""
    document = docx.Document(filepath)
    text_parts = [para.text for para in document.paragraphs if para.text.strip()]

    # Also pull text out of any tables in the document (some resumes use
    # tables for layout, e.g. skills or contact info sections)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    return "\n".join(text_parts).strip()


def extract_text(filepath, file_extension):
    """
    Dispatch to the correct extractor based on file extension.
    Returns extracted text, or raises ValueError for unsupported types.
    """
    file_extension = file_extension.lower()
    if file_extension == "pdf":
        return extract_text_from_pdf(filepath)
    elif file_extension == "docx":
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")
